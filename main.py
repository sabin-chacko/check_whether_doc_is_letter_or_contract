import os
from xmlrpc.client import Boolean
import docx
import requests
from urllib.parse import quote
#location of docx files
data_folder = "data_folder"

# contract keywords
contract_keywords = ["AGREEMENT", "TERM", "CONFIDENTIALITY", "NONCOMPETE", "CONTRACT", "SIGNATURE"]
# letter keywords
letter_keywords = ["THIS LETTER", "SINCERELY,", "RESPECTFULLY YOURS,", "WRITING TO INFORM", "YOURS FAITHFULLY"]

def get_docx_filename() -> list:
    """to get docx file from the folder

    Returns:
        list: list of docx file name
    """    
    file_list = []

    if os.path.exists(data_folder):
        for f in os.listdir(data_folder):
            #check f is of type file and f is of xlsx format
            if os.path.isfile(os.path.join(data_folder, f)) and\
                f.split(".")[1] == "docx":
                file_list.append(f)
    return file_list
    
def get_bold_list(bold_list, para):
    for run in para.runs:
        if run.bold:
            bold_list.append(run.text)
    return bold_list

def get_bold_text(doc):
    bold_list= []
    for para in doc.paragraphs:
        bold_list = get_bold_list(bold_list, para)
    return bold_list

def is_contract(doc: object) -> bool:
    """check whether the doc is contract by traversing the bold heading 
       with the contract keywords

    Args:
        doc (object): document

    Returns:
        bool: whether doc is contract or not
    """    
    try:
        # get each bold text
        bold_list = get_bold_text(doc)
        for keys in bold_list:
            for val in contract_keywords:
                #check if bold words contain contract keywords
                if keys.upper().find(val) > 0:
                    return True
        return False
    except Exception as e:
        print(str(e))
        False

def is_letter(doc: object) -> bool:
    """check whether the doc is letter by traversing each line
       with the letter keywords

    Args:
        doc (object): document

    Returns:
        bool: whether doc is letter or not
    """
    try:
        # check if each paragraph contains letter keywords
        for para in doc.paragraphs:
            for keys in letter_keywords:
                if para.text.upper().find(keys) > 0:
                    return True
        return False
    except Exception as e:
        print(str(e))
        return False


def proccess_each_doc(filename: str) -> None:
    """process each file

    Args:
        filename (str): filename

    """
    # initialise address string
    headings = []
    # read the document
    doc = docx.Document(os.path.join(data_folder, filename))
    
    if is_contract(doc):
        print(f'{filename} is Contract')
    elif is_letter(doc):
        print(f'{filename} is Letter')
    else:
        print(f'{filename} not identified')

def main():
    """main method
    """
    file_list = get_docx_filename()
    for f in file_list:
        proccess_each_doc(f)


if __name__ == "__main__":
    main()